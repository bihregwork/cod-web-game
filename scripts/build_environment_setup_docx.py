from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ENVIRONMENT_SETUP.md"
OUT = ROOT / "docs" / "ENVIRONMENT_SETUP.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(8.8)
    if bold:
        run.font.color.rgb = RGBColor(30, 55, 86)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    body = rows[2:] if len(rows) > 1 and all(set(c.replace(" ", "")) <= {"-"} for c in rows[1]) else rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "E8EEF5")
    for row in body:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(headers)]):
            set_cell_text(cells[index], value.replace("`", ""))
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def add_inline_code(paragraph, text: str) -> None:
    for index, part in enumerate(text.split("`")):
        run = paragraph.add_run(part)
        if index % 2 == 1:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(30, 55, 86)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_code(p, line[2:].strip())
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline_code(p, line)
        index += 1

    footer = section.footer.paragraphs[0]
    footer.text = "Cod_Web_Game | Environment Setup"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
