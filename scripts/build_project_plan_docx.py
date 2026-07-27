from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PROJECT_PLAN.md"
OUT = ROOT / "docs" / "PROJECT_PLAN.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(8.3)
    if bold:
        run.font.color.rgb = RGBColor(30, 55, 86)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    separator_idx = 1 if len(rows) > 1 and all(set(c.replace(" ", "")) <= {"-"} for c in rows[1]) else None
    body = rows[2:] if separator_idx == 1 else rows[1:]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip().strip("|")
        rows.append([cell.strip().replace("`", "") for cell in line.split("|")])
        idx += 1
    return rows, idx


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(35, 35, 35)


def add_text_with_inline_code(paragraph, text: str) -> None:
    parts = text.split("`")
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        if line.startswith("```"):
            code_lines: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(doc, code_lines)
            idx += 1
            continue
        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(30, 55, 86)
            idx += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            idx += 1
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_text_with_inline_code(paragraph, line[2:].strip())
            idx += 1
            continue
        paragraph = doc.add_paragraph()
        add_text_with_inline_code(paragraph, line)
        idx += 1

    footer = section.footer.paragraphs[0]
    footer.text = "Cod_Web_Game | Project Plan"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
